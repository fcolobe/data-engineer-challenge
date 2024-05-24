import logging
import os
import re
from datetime import datetime

import pandas as pd
import pymupdf
from docx import Document
from lxml import etree

from exo_1 import insert_new_data

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)


def get_pdf_and_docx_files(directory):
    """
    Returns the names of .pdf and .docx files in the specified directory.

    Args:
    - directory (str): The path to the directory.

    Returns:
    - list: A list containing the names of .pdf and .docx files.
    """
    pdf_and_docx_files = []
    try:
        for filename in os.listdir(directory):
            if filename.endswith(".pdf") or filename.endswith(".docx"):
                pdf_and_docx_files.append(filename)
    except Exception as e:
        logging.error(f"Error reading the directory: {e}")
    return pdf_and_docx_files


def extract_text_from_pdf(filepath):
    """
    Extracts text from a PDF file.

    Args:
    - filepath (str): The path to the PDF file.

    Returns:
    - str: The extracted text.
    """
    try:
        doc = pymupdf.open(filepath)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    except Exception as e:
        logging.error(f"Error reading the PDF: {e}")
        return ""


def extract_textboxes_from_docx(doc):
    """
    Extracts text from text boxes in a DOCX file.

    Args:
    - doc: The Document object representing the DOCX file.

    Returns:
    - str: The extracted text.
    """
    text = []
    seen_texts = set()

    doc_xml = doc.element.xml
    root = etree.fromstring(doc_xml.encode("utf-8"))

    for shape in root.xpath(
        "//w:txbxContent",
        namespaces={
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        },
    ):
        for paragraph in shape.xpath(
            ".//w:p",
            namespaces={
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            },
        ):
            text_elements = paragraph.xpath(
                ".//w:t",
                namespaces={
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                },
            )
            for t in text_elements:
                if t.text:
                    if t.text and t.text not in seen_texts:
                        seen_texts.add(t.text)
                        text.append(t.text)

    return "\n".join(text)


def extract_text_from_docx(filepath):
    """
    Extracts text from a DOCX file, including paragraphs, tables, and text boxes.

    Args:
    - filepath (str): The path to the DOCX file.

    Returns:
    - str: The extracted text.
    """
    try:
        doc = Document(filepath)
        text = []

        textbox_text = extract_textboxes_from_docx(doc)
        if textbox_text:
            text.append(textbox_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text.append(cell_text)

        for para in doc.paragraphs:
            text.append(para.text)

        return "\n".join(text)
    except Exception as e:
        logging.error(f"Error reading the DOCX: {e}")
        return ""


def normalize_text(text):
    """
    Normalizes the text.

    Args:
    - text (str): The text to normalize.

    Returns:
    - str: The normalized text.
    """
    text = text.strip()
    text = re.sub(r"\s+", " ", text)
    text = text.lower()
    return text


def capitalize_author(prefix, author):
    """
    Capitalize the author's name and include the prefix.

    Args:
    - prefix (str): The prefix of the author's name.
    - author (str): The author's name.

    Returns:
    - str: The capitalized author's name including the prefix.
    """
    return f"{prefix.capitalize()} {' '.join(word.capitalize() for word in author.split())}"


def extract_metadata(text):
    """
    Extracts metadata such as the date and author from the document text.

    Args:
    - text (str): The text of the document.

    Returns:
    - tuple: A tuple containing the document date and author extracted from the text.
    """
    if not text:
        return None, None

    normalized_text = normalize_text(text)

    date_pattern = r"\b\d{2}/\d{2}/\d{4}\b"
    dates = re.findall(date_pattern, normalized_text)
    document_date = None

    for date_str in dates:
        date_obj = datetime.strptime(date_str, "%d/%m/%Y")
        if date_obj.year >= 2001:
            document_date = datetime.strftime(date_obj, "%d/%m/%Y")
            break

    author_pattern = r"\b(dr)\s+([a-z]+(?:\s+[a-z]+)?)\b"
    author_matches = re.findall(author_pattern, normalized_text)
    if author_matches:
        last_author = author_matches[-1]
        prefix, full_name = last_author
        full_name = full_name.split("dr")[0]
        author = capitalize_author(prefix, full_name)
    else:
        author = None

    return document_date, author


def get_patient_info(conn, ipp):
    """
    Retrieves patient information from the DWH_PATIENT_IPPHIST table.

    Args:
    - conn (sqlite3.Connection): The connection to the database.
    - ipp (str): The IPP (Identifiant Patient Principal) of the patient.

    Returns:
    - tuple: A tuple containing the patient information retrieved from the database.
    """
    query = "SELECT * FROM DWH_PATIENT_IPPHIST WHERE HOSPITAL_PATIENT_ID = ?"
    cursor = conn.execute(query, (ipp,))
    return cursor.fetchone()


def get_document_data(directory, upload_id, conn):
    """
    Processes PDF and DOCX files in a directory and extracts their metadata.

    Args:
    - directory (str): The path to the directory containing the files.
    - upload_id (int): The upload ID to associate with the records.
    - conn (sqlite3.Connection): The connection to the database.

    Returns:
    - list: A list of dictionaries containing document data.
    """
    files = get_pdf_and_docx_files(directory)
    document_num = 1
    documents = []

    for file in files:
        file_path = os.path.join(directory, file)
        file_name, file_extension = os.path.splitext(file)
        ipp, id_document = file_name.split("_")
        displayed_text = ""

        if file_extension == ".pdf":
            displayed_text = extract_text_from_pdf(file_path)
            document_origin_code = "DOSSIER_PATIENT"
        elif file_extension == ".docx":
            displayed_text = extract_text_from_docx(file_path)
            document_origin_code = "RADIOLOGIE_SOFTWARE"

        if not displayed_text:
            logging.warning(f"The file {file} is empty or could not be read.")
            continue

        document_date, author = extract_metadata(displayed_text)

        patient_info = get_patient_info(conn, ipp)
        if not patient_info:
            logging.warning(f"No information found for the patient: {ipp}")
            continue

        document_dict = {
            "DOCUMENT_NUM": document_num,
            "PATIENT_NUM": patient_info[0],
            "ENCOUNTER_NUM": None,
            "TITLE": None,
            "DOCUMENT_ORIGIN_CODE": document_origin_code,
            "DOCUMENT_DATE": document_date,
            "ID_DOC_SOURCE": id_document,
            "DOCUMENT_TYPE": file_extension.lstrip("."),
            "DISPLAYED_TEXT": displayed_text,
            "AUTHOR": author,
            "UNIT_CODE": None,
            "UNIT_NUM": None,
            "DEPARTMENT_NUM": None,
            "EXTRACTCONTEXT_DONE_FLAG": 0,
            "EXTRACTCONCEPT_DONE_FLAG": 0,
            "ENRGENE_DONE_FLAG": 0,
            "ENRICHTEXT_DONE_FLAG": 0,
            "UPLOAD_ID": upload_id,
        }

        documents.append(document_dict)
        document_num += 1

    return documents


def update_existing_doc_data(df, table_name, conn):
    """
    Updates existing records in a database table with data from a DataFrame.

    Args:
    - df (pd.DataFrame): The DataFrame containing the data to update.
    - table_name (str): The name of the target table.
    - conn (sqlite3.Connection): The connection to the database.

    Returns:
    - None
    """
    for _, row in df.iterrows():
        placeholders = ", ".join(
            [f"{col} = ?" for col in row.index if col != "DOCUMENT_NUM"]
        )
        query = f"UPDATE {table_name} SET {placeholders} WHERE DOCUMENT_NUM = ?"
        params = tuple(row[col] for col in row.index if col != "DOCUMENT_NUM") + (
            row["DOCUMENT_NUM"],
        )
        conn.execute(query, params)
    conn.commit()


def update_document_data(directory, upload_id, conn):
    """
    Updates the DWH_DOCUMENT table with data extracted from PDF and DOCX files.

    Args:
    - directory (str): The path to the directory containing the files.
    - upload_id (int): The upload ID to associate with the records.
    - conn (sqlite3.Connection): The connection to the database.

    Returns:
    - None
    """
    try:
        document_data = get_document_data(directory, upload_id, conn)
        df_documents = pd.DataFrame(document_data)

        existing_documents = pd.read_sql_query("SELECT * FROM DWH_DOCUMENT", conn)

        new_documents = df_documents[
            ~df_documents["DOCUMENT_NUM"].isin(existing_documents["DOCUMENT_NUM"])
        ]

        update_existing_doc_data(df_documents, "DWH_DOCUMENT", conn)
        insert_new_data(new_documents, "DWH_DOCUMENT", conn)

        logging.info("Update successful.")

    except Exception as e:
        logging.error(f"Error updating DWH_DOCUMENT table: {e}")
